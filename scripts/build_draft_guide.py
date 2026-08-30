"""Render the printable 2026 Draft Guide docx for a given draft slot.

Reads the finished big board (data/derived/big_board_PPR.parquet) so every
"My Rank", projection and Value figure in the guide is the model's own number,
then lays out a round-by-round page of player cards.

    python scripts/build_draft_guide.py --slot 5 --out ~/Desktop/2026_Draft_Guide_Brandon.docx

Card content (which player is THE PICK, the alternates, the AVOID call and the
prose) is hand-authored in ROUNDS below — the script supplies the numbers and
the artwork so the two can't drift apart.
"""
from __future__ import annotations

import argparse
import hashlib
import io
import re
import urllib.request
from pathlib import Path

import pandas as pd
from PIL import Image
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
BOARD = ROOT / "data" / "derived" / "big_board_PPR.parquet"
CACHE = ROOT / "data" / "derived" / "guide_art"

# ── Palette ──────────────────────────────────────────────────────────────────
NAVY = "1B2A4A"
NAVY_TINT = "9FB3D1"
INK = "1A1A1A"
MUTED = "5A5F67"
FAINT = "8A8F98"
RULE = "C9CED6"
GREEN = "1A6B45"
GREEN_BG = "F1F6F3"
RED = "A3281E"
RED_BG = "FBF0EF"
GREY_BG = "F7F8FA"

TEAMS = 11
ROUNDS_SHOWN = 10


# ── Board lookup ─────────────────────────────────────────────────────────────
def load_board() -> pd.DataFrame:
    b = pd.read_parquet(BOARD).sort_values("vor", ascending=False).reset_index(drop=True)
    b["my_rank"] = b.index + 1
    return b


def pick_numbers(slot: int, rounds: int = ROUNDS_SHOWN) -> list[int]:
    """Snake-draft overall pick numbers for a slot."""
    return [
        rd * TEAMS + slot if rd % 2 == 0 else rd * TEAMS + (TEAMS - slot + 1)
        for rd in range(rounds)
    ]


def stats(board: pd.DataFrame, name: str) -> dict:
    row = board[board.player == name]
    if row.empty:
        raise SystemExit(f"'{name}' is not on the big board — check the spelling.")
    r = row.iloc[0]
    espn = r.espn_overall
    return {
        "name": name,
        "pos": r.pos,
        "team": r.team,
        "my_rank": int(r.my_rank),
        "espn": "—" if pd.isna(espn) else f"#{int(espn)}",
        "pts": f"{r.predicted_pts:.1f}",
        "ppg": f"{r.pred_ppg:.1f}",
        "games": f"{r.proj_games:.1f}",
        "vor": f"{r.vor:+.1f}",
    }


# ── Artwork ──────────────────────────────────────────────────────────────────
def _fetch(url: str, tag: str, box: int) -> Path | None:
    """Download, square-crop and downscale to `box` px.

    The raw ESPN headshots are 3400x2450 and ~5MB each; embedded as-is they
    push the finished document past 100MB. Cropping to a square keeps the
    circular-looking card art and resizing to the size it actually renders at
    puts the file back near 3MB.
    """
    CACHE.mkdir(parents=True, exist_ok=True)
    dest = CACHE / f"{tag}_{hashlib.md5(url.encode()).hexdigest()[:10]}_{box}.png"
    if dest.exists():
        return dest
    try:
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        data = urllib.request.urlopen(req, timeout=20).read()
    except Exception:
        return None

    im = Image.open(io.BytesIO(data)).convert("RGBA")
    w, h = im.size
    side = min(w, h)
    # Horizontally centred; vertically biased to the top so heads aren't cropped.
    left = (w - side) // 2
    top = 0 if h > w else (h - side) // 2
    im = im.crop((left, top, left + side, top + side))
    im = im.resize((box, box), Image.LANCZOS)
    im.save(dest, "PNG", optimize=True)
    return dest


def _normalize_name(name: str) -> str:
    """Match utils.nfl_data_core.normalize_name — suffix/punctuation-insensitive."""
    name = str(name).lower().strip()
    name = re.sub(r"\s+(jr\.?|sr\.?|ii|iii|iv)$", "", name)
    name = re.sub(r"[.\-']", "", name)
    return re.sub(r"\s+", " ", name).strip()


def art_maps() -> tuple[dict, dict]:
    wk = pd.read_csv(
        ROOT / "data" / "raw" / "weekly.csv",
        low_memory=False,
        usecols=["player_display_name", "headshot_url", "season"],
    )
    shots = {
        _normalize_name(k): v for k, v in
        wk.sort_values("season")
        .groupby("player_display_name")["headshot_url"]
        .last()
        .dropna()
        .items()
    }
    # Current-season roster pull wins: weekly.csv stops at the last completed
    # season, so on its own the guide prints last year's uniform for anyone who
    # moved and nothing at all for rookies.
    cur = ROOT / "data" / "raw" / "headshots.csv"
    if cur.exists():
        h = pd.read_csv(cur).dropna(subset=["player_name", "headshot_url"])
        shots.update(zip(h["player_name"].map(_normalize_name), h["headshot_url"]))

    tm = pd.read_csv(ROOT / "data" / "raw" / "teams.csv")
    logos = dict(zip(tm.team_abbr, tm.team_logo_espn))
    return shots, logos


# ── Low-level docx helpers ───────────────────────────────────────────────────
def shade(cell, hex_fill: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def no_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "none")
        e.set(qn("w:sz"), "0")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def cell_margins(table, top=70, left=0, bottom=70, right=110) -> None:
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    table._tbl.tblPr.append(mar)


def bottom_rule(par, color=RULE) -> None:
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), "6")
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color)
    pbdr.append(b)
    par._p.get_or_add_pPr().append(pbdr)


def run(par, text, *, size=20, bold=False, italic=False, color=INK,
        caps=False, track=None):
    r = par.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(size / 2)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = RGBColor.from_string(color)
    rpr = r._element.get_or_add_rPr()
    if caps:
        rpr.append(OxmlElement("w:caps"))
    if track:
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), str(track))
        rpr.append(sp)
    return r


def para(container, *, before=None, after=0, align=None, line=None):
    """Spacing args are in twips (as they appear in the XML): 20 twips = 1pt."""
    p = container.add_paragraph()
    pf = p.paragraph_format
    if before is not None:
        pf.space_before = Pt(before / 20)
    pf.space_after = Pt(after / 20)
    if align:
        p.alignment = align
    if line:
        pfmt = p._p.get_or_add_pPr()
        sp = pfmt.find(qn("w:spacing"))
        if sp is None:
            sp = OxmlElement("w:spacing")
            pfmt.append(sp)
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "auto")
    return p


def set_widths(table, widths: list[int]) -> None:
    """Widths are dxa (twentieths of a point). Both the grid and the cells have
    to carry them — renderers disagree about which one wins under fixed layout."""
    table.autofit = False
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table._tbl.tblPr.append(layout)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, w in zip(grid.findall(qn("w:gridCol")), widths):
            col.set(qn("w:w"), str(w))

    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Emu(w * 635)


# ── Blocks ───────────────────────────────────────────────────────────────────
def cover(doc, slot: int, picks: list[int]) -> None:
    for _ in range(4):
        para(doc, before=0, after=0)

    p = para(doc, before=0, after=40, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "2026 FANTASY FOOTBALL", size=26, bold=True, color=FAINT, track=60)

    p = para(doc, before=0, after=120, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Draft Guide", size=80, bold=True, color=NAVY)

    p = para(doc, before=0, after=480, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, f"Round-by-round targets for pick #{slot}", size=26, italic=True, color=MUTED)
    bottom_rule(p)

    meta = [
        ("League   ", f"{TEAMS} teams  ·  ESPN"),
        ("Draft slot   ", f"Pick #{slot}"),
        ("Scoring   ", "Full PPR"),
        ("Your picks   ", "  ·  ".join(str(n) for n in picks)),
    ]
    for label, value in meta:
        p = para(doc, before=0, after=60, align=WD_ALIGN_PARAGRAPH.CENTER)
        run(p, label, size=20, bold=True, color=FAINT, caps=True)
        run(p, value, size=22, color=INK)

    p = para(doc, before=240, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Every ranking marked \u201cMy Rank\u201d comes from your own projection model.",
        size=20, italic=True, color=MUTED)
    p = para(doc, before=0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    run(p, "Where it disagrees with ESPN, that gap is the edge.",
        size=20, italic=True, color=MUTED)


def rules(doc, items: list[tuple[str, str]]) -> None:
    doc.add_page_break()
    p = para(doc, before=0, after=200)
    run(p, "The Three Rules", size=44, bold=True, color=NAVY)
    bottom_rule(p)

    for i, (head, body) in enumerate(items, start=1):
        p = para(doc, before=200, after=60)
        run(p, f"{i}.  ", size=28, bold=True, color=GREEN)
        run(p, head, size=28, bold=True, color=NAVY)
        p = para(doc, before=0, after=0, line=269)
        run(p, body, size=21, color=INK)


def cheat_sheet(doc, board, rounds, picks) -> None:
    # Shares a page with the Three Rules — no break here.
    p = para(doc, before=400, after=200)
    run(p, "Cheat Sheet", size=44, bold=True, color=NAVY)
    bottom_rule(p)

    t = doc.add_table(rows=1 + len(rounds), cols=4)
    no_borders(t)
    cell_margins(t, top=90, left=120, bottom=90, right=120)
    widths = [900, 900, 3800, 3760]

    hdr = ["ROUND", "PICK", "THE PICK", "IF HE'S GONE"]
    for c, text in zip(t.rows[0].cells, hdr):
        shade(c, NAVY)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        pp = c.paragraphs[0]
        pp.paragraph_format.space_after = Pt(0)
        run(pp, text, size=18, bold=True, color="FFFFFF", track=40)

    for i, (rd, pick) in enumerate(zip(rounds, picks), start=1):
        cells = t.rows[i].cells
        for c in cells:
            shade(c, "FFFFFF" if i % 2 else GREY_BG)
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.paragraphs[0].paragraph_format.space_after = Pt(0)
        run(cells[0].paragraphs[0], str(rd["round"]), size=20, bold=True, color=NAVY)
        run(cells[1].paragraphs[0], f"#{pick}", size=20, color=MUTED)
        main = rd["cards"][0]
        run(cells[2].paragraphs[0], main["player"], size=20, bold=True, color=INK)
        alt = rd["cards"][1]["player"] if len(rd["cards"]) > 1 else "—"
        run(cells[3].paragraphs[0], alt, size=20, color=MUTED)
    set_widths(t, widths)

    p = para(doc, before=200, after=0)
    run(p, "Rounds 11+: defense and kicker. Nothing else matters that late.",
        size=20, italic=True, color=MUTED)


def round_header(doc, rd: int, pick: int, title: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    no_borders(t)
    cell_margins(t, top=110, left=150, bottom=110, right=150)
    c = t.rows[0].cells[0]
    shade(c, NAVY)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run(p, f"ROUND {rd}", size=22, bold=True, color=NAVY_TINT, track=60)
    run(p, f"     PICK #{pick}", size=22, bold=True, color="FFFFFF", track=40)
    p2 = para(c, before=0, after=0)
    run(p2, title, size=30, bold=True, color="FFFFFF")
    set_widths(t, [9360])


def player_card(doc, s: dict, badge: str, kind: str, blurb: str, shots, logos) -> None:
    fill = {"pick": GREEN_BG, "alt": GREY_BG, "avoid": RED_BG}[kind]
    bcol = {"pick": GREEN, "alt": MUTED, "avoid": RED}[kind]

    t = doc.add_table(rows=1, cols=3)
    no_borders(t)
    cell_margins(t)
    cells = t.rows[0].cells
    for c in cells:
        shade(c, fill)
    cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cells[2].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    # headshot
    p = cells[0].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    url = shots.get(_normalize_name(s["name"]))
    path = _fetch(url, "h", 340) if url else None
    if path:
        p.add_run().add_picture(str(path), width=Emu(658368))

    # team logo
    p = cells[1].paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lurl = logos.get(s["team"])
    lpath = _fetch(lurl, f"t{s['team']}", 200) if lurl else None
    if lpath:
        p.add_run().add_picture(str(lpath), width=Emu(310896))

    # body
    body = cells[2]
    p = body.paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    run(p, s["name"], size=26, bold=True, color=NAVY)
    run(p, f"   {s['pos']} · {s['team']}", size=19, bold=True, color=FAINT)

    p = para(body, before=0, after=60)
    run(p, badge, size=16, bold=True, color=bcol, track=40)

    p = para(body, before=0, after=80)
    fields = [
        ("My rank ", f"#{s['my_rank']}"),
        ("ESPN ", s["espn"]),
        ("Proj pts ", s["pts"]),
        ("Per game ", s["ppg"]),
        ("Value ", s["vor"]),
    ]
    for i, (label, value) in enumerate(fields):
        if i:
            run(p, "   |   ", size=18, color=RULE)
        run(p, label, size=17, color=FAINT, caps=True)
        run(p, value, size=20, bold=True, color=INK)

    p = para(body, before=0, after=0, line=269)
    run(p, blurb, size=20, color=INK)

    set_widths(t, [1150, 620, 7590])
    para(doc, before=0, after=80)


def build(slot: int, out: Path, rounds_spec: list[dict]) -> None:
    board = load_board()
    shots, logos = art_maps()
    picks = pick_numbers(slot)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Emu(7772400), Emu(10058400)  # US Letter
    sec.top_margin, sec.bottom_margin = Emu(731520), Emu(640080)
    sec.left_margin = sec.right_margin = Emu(914400)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    cover(doc, slot, picks)
    rules(doc, RULES)
    cheat_sheet(doc, board, rounds_spec, picks)

    for rd, pick in zip(rounds_spec, picks):
        doc.add_page_break()
        round_header(doc, rd["round"], pick, rd["title"])
        para(doc, before=0, after=60)
        p = para(doc, before=0, after=140, line=269)
        run(p, rd["intro"], size=21, color=INK)

        for card in rd["cards"]:
            player_card(doc, stats(board, card["player"]), card["badge"],
                        card["kind"], card["blurb"], shots, logos)

        avoid = rd.get("avoid", [])
        if avoid:
            p = para(doc, before=120, after=60)
            run(p, "DO NOT DRAFT", size=20, bold=True, color=RED, track=60)
            for card in avoid:
                player_card(doc, stats(board, card["player"]), card["badge"],
                            "avoid", card["blurb"], shots, logos)

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"wrote {out}")


# ── Content ──────────────────────────────────────────────────────────────────
RULES = [
    ("Take Jaxon Smith-Njigba at pick 5.",
     "The four names ESPN ranks above him — Gibbs, Chase, Nacua, Bijan — will all be "
     "gone. He is your model's #4 overall, projects 335.7 points across 15 games, and "
     "carries no injury flag. ESPN has him 5th too, so nobody is reaching past him to "
     "take him first. He should simply be there."),
    ("Take a tight end in round 2.",
     "There are two useful tight ends this year and then a 45-point cliff. Trey McBride "
     "and Brock Bowers are the only ones that matter — McBride will be gone by pick 18, "
     "so Bowers is the one. Miss both and you are starting a tight end worth 45 fewer "
     "points than the guy your league-mate got."),
    ("Wait on quarterback — unless Lamar falls to round 4.",
     "After the top three, every quarterback from QB4 through QB13 projects within 15 "
     "points of one another, so there is no reason to spend a real pick there. The one "
     "exception: your model has Lamar Jackson 32nd overall and ESPN has him 39th. If he "
     "is on the board at pick 40, take him and never think about the position again."),
]

ROUNDS_SPEC = [
    {
        "round": 1, "title": "Take the receiver, not the running back",
        "intro": "The top four are gone. ESPN's board will show Christian McCaffrey as "
                 "the best player left — ignore it. Smith-Njigba is worth 100 more "
                 "points than McCaffrey in your model.",
        "cards": [
            {"player": "Jaxon Smith-Njigba", "badge": "THE PICK", "kind": "pick",
             "blurb": "Your model's #4 overall and the clear best player available at 5. "
                      "335.7 projected points over 15.1 games with no injury flag — the "
                      "only receiver in this tier who is both elite and healthy. ESPN "
                      "agrees enough (#5) that he will not be reached for ahead of you."},
            {"player": "Amon-Ra St. Brown", "badge": "IF SOMEONE TAKES HIM", "kind": "alt",
             "blurb": "Essentially the same pick one notch down, and the player your wife "
                      "is hoping falls to 6. Team-leading target share in the best "
                      "receiving offense in football, 14.9 projected games, no injury flag."},
            {"player": "Jonathan Taylor", "badge": "IF YOU WANT THE RB", "kind": "alt",
             "blurb": "The only running back worth taking here. 294.5 points and 20.7 a "
                      "game, but he carries an injury flag and 14.2 projected games. Take "
                      "him only if both receivers are gone."},
        ],
        "avoid": [
            {"player": "Christian McCaffrey", "badge": "AVOID",
             "blurb": "ESPN's #6 and the trap of the first round. Your model has him 17th "
                      "— age-29 curve coming off an Achilles, only 13.6 projected games. "
                      "He is worth about half of Smith-Njigba. Let someone else take him."},
        ],
    },
    {
        "round": 2, "title": "Grab the second of the two tight ends",
        "intro": "McBride will be gone. Bowers is the last tight end who separates from "
                 "the field — after him the position falls off a 45-point cliff and every "
                 "option looks the same until round 12.",
        "cards": [
            {"player": "Brock Bowers", "badge": "THE PICK", "kind": "pick",
             "blurb": "16 projected games, no injury flag, and 240 points at a position "
                      "where the 12th-best option projects 155. That 45-point gap is the "
                      "single largest positional edge available to you all draft."},
            {"player": "George Pickens", "badge": "IF BOWERS IS GONE", "kind": "alt",
             "blurb": "Your model's #11 overall and ESPN's #29 — an 18-slot discount. If "
                      "Bowers is taken, this is the value play, and Pickens has a real "
                      "chance to still be there at 27 anyway."},
            {"player": "Omarion Hampton", "badge": "FALLBACK", "kind": "alt",
             "blurb": "16 projected games and a clean bill of health at running back, "
                      "which is rare in this range. Model #14, ESPN #32."},
        ],
        "avoid": [
            {"player": "Saquon Barkley", "badge": "AVOID",
             "blurb": "ESPN #21, your model #49. The gap is 28 slots — the largest reach "
                      "on the board in this range. 190.7 projected points is RB-committee "
                      "territory, not a second-round price."},
            {"player": "Garrett Wilson", "badge": "AVOID",
             "blurb": "ESPN #22, your model #48. Only 13.2 projected games and a target "
                      "share that never recovered. There are better receivers 20 picks later."},
        ],
    },
    {
        "round": 3, "title": "Cash in the biggest discount on the board",
        "intro": "Pickens is the play if he lasted. Both he and Hampton are top-14 in "
                 "your model but priced by ESPN in the 30s, and there is a real chance "
                 "both are sitting here.",
        "cards": [
            {"player": "George Pickens", "badge": "THE PICK", "kind": "pick",
             "blurb": "265.1 points, 17.9 a game, no injury flag — WR numbers that belong "
                      "in round 2. ESPN's #29 ranking is the single biggest mispricing in "
                      "the top 30 of your board."},
            {"player": "Omarion Hampton", "badge": "NEARLY AS GOOD", "kind": "alt",
             "blurb": "Model #14 against an ESPN #32 price. A full 16-game projection at "
                      "running back with no flag is worth more than the raw points suggest."},
            {"player": "Malik Nabers", "badge": "THIRD OPTION", "kind": "alt",
             "blurb": "240.6 points and 16.0 a game, but the injury flag is real — coming "
                      "off the ACL that cost him 2025. Take him only if the first two are gone."},
        ],
        "avoid": [
            {"player": "Jayden Daniels", "badge": "AVOID",
             "blurb": "ESPN #34, your model #64. Spending a third-round pick on a "
                      "quarterback worth +12.0 over replacement when Kyler Murray at pick "
                      "93 is worth +1.5 is the exact mistake rule 3 exists to prevent."},
        ],
    },
    {
        "round": 4, "title": "Start the run on discounted running backs",
        "intro": "This is where your board and ESPN's diverge hardest at running back. "
                 "Everything from here through round 6 is buying at a discount — and if "
                 "Lamar Jackson slipped this far, he is the exception worth taking.",
        "cards": [
            {"player": "Cam Skattebo", "badge": "THE PICK", "kind": "pick",
             "blurb": "Model #19 against an ESPN #56 price — a 37-slot discount, the "
                      "largest of the entire middle rounds. 16 projected games, no injury "
                      "flag, 228.5 points."},
            {"player": "Lamar Jackson", "badge": "ONLY QB WORTH IT", "kind": "alt",
             "blurb": "The single exception to waiting on quarterback. Model #32, ESPN "
                      "#39, and +43.2 over replacement — more than double the next tier. "
                      "If he is here at 40, take him."},
            {"player": "Kyren Williams", "badge": "IF SKATTEBO IS GONE", "kind": "alt",
             "blurb": "226.0 points and 15.5 a game at an ESPN #49 price. Carries an "
                      "injury flag and 14.6 projected games, but the volume is not in doubt."},
        ],
        "avoid": [
            {"player": "Davante Adams", "badge": "AVOID",
             "blurb": "ESPN #45, your model #76. 170.2 projected points is replacement "
                      "level at receiver — you can get the same production in round 9."},
        ],
    },
    {
        "round": 5, "title": "The best value left is still a running back",
        "intro": "Two backs your model rates inside the top 22 overall are both priced "
                 "here. If either is available, do not overthink it.",
        "cards": [
            {"player": "D'Andre Swift", "badge": "THE PICK", "kind": "pick",
             "blurb": "Model #22 against ESPN #66 — a 44-slot discount and the last "
                      "genuine steal of the draft. 223.5 points and 15.1 a game as the "
                      "clear lead back in Chicago."},
            {"player": "Travis Etienne", "badge": "JUST AS GOOD", "kind": "alt",
             "blurb": "220.5 points in the New Orleans backfield at an ESPN #46 price. "
                      "Model #25 — the gap to Swift is inside the margin of error."},
            {"player": "Jaylen Waddle", "badge": "IF YOU NEED WR", "kind": "alt",
             "blurb": "205.9 points after the move to Denver, priced at ESPN #58 against "
                      "a model rank of #36. The best receiver available in this range."},
        ],
        "avoid": [
            {"player": "Bhayshul Tuten", "badge": "AVOID",
             "blurb": "ESPN #62, your model #132 — a 70-slot gap and negative value over "
                      "replacement. The worst price-to-projection mismatch on the board."},
        ],
    },
    {
        "round": 6, "title": "One more back, then start filling receiver",
        "intro": "Swift may still be here — ESPN prices him at #66, four picks after you. "
                 "If he is gone, Alec Pierce is the largest remaining gap between your "
                 "model and the field.",
        "cards": [
            {"player": "Alec Pierce", "badge": "THE PICK", "kind": "pick",
             "blurb": "Model #42, ESPN #92 — a 50-slot discount, the biggest left on the "
                      "board. 197.2 points across a full 15 games with no injury flag."},
            {"player": "Jaylen Warren", "badge": "GOOD BACKUP PLAN", "kind": "alt",
             "blurb": "189.4 points as the lead back in Pittsburgh at an ESPN #84 price. "
                      "Model #51."},
            {"player": "George Kittle", "badge": "TE INSURANCE", "kind": "alt",
             "blurb": "Only worth it if you missed Bowers. 13.0 projected games coming off "
                      "the Achilles, but 13.5 a game is still top-8 at the position."},
        ],
        "avoid": [
            {"player": "David Montgomery", "badge": "AVOID",
             "blurb": "ESPN #69, your model #140. Negative value over replacement in a "
                      "backfield he no longer leads."},
        ],
    },
    {
        "round": 7, "title": "Depth, and a second tight end if you missed",
        "intro": "The bargains are mostly gone. From here, target players your model "
                 "still rates well above their price and fill out the bench.",
        "cards": [
            {"player": "Alec Pierce", "badge": "THE PICK", "kind": "pick",
             "blurb": "If he somehow lasted this long he is the easiest pick of the "
                      "middle rounds. ESPN does not have him inside their top 90; your "
                      "model has him 42nd."},
            {"player": "Jake Ferguson", "badge": "TE INSURANCE", "kind": "alt",
             "blurb": "15.7 projected games — the most of any tight end in this range. "
                      "Only relevant if you missed Bowers and Kittle."},
            {"player": "TreVeyon Henderson", "badge": "UPSIDE BACK", "kind": "alt",
             "blurb": "179.5 points with real weekly upside if the New England backfield "
                      "consolidates. Model #62, ESPN #76."},
        ],
    },
    {
        "round": 8, "title": "Take the last of the real value",
        "intro": "Jaylen Warren is the priority if he made it this far. After this pick "
                 "the board flattens and everything is worth roughly the same.",
        "cards": [
            {"player": "Jaylen Warren", "badge": "THE PICK", "kind": "pick",
             "blurb": "189.4 points and +23.9 over replacement is startable production at "
                      "a bench price. The last player on the board worth more than 20 "
                      "points over replacement."},
            {"player": "Brian Thomas Jr.", "badge": "STRONG BACKUP", "kind": "alt",
             "blurb": "Model #67, ESPN #106. 174.3 points with the target share to beat "
                      "that if the Jacksonville offense takes a step."},
            {"player": "Parker Washington", "badge": "NAME UPSIDE", "kind": "alt",
             "blurb": "172.1 points and 12.8 a game — quietly the WR2 in Jacksonville at "
                      "an ESPN #88 price."},
        ],
        "avoid": [
            {"player": "Mark Andrews", "badge": "AVOID",
             "blurb": "ESPN #90, your model #133. Negative value over replacement — the "
                      "Baltimore target share moved on without him."},
        ],
    },
    {
        "round": 9, "title": "Now, finally, take your quarterback",
        "intro": "This is the whole point of waiting. If you did not get Lamar in round "
                 "4, every quarterback from here to round 12 projects within a few points "
                 "of one another — so take one and move on.",
        "cards": [
            {"player": "Kyler Murray", "badge": "THE PICK", "kind": "pick",
             "blurb": "250.7 points and 17.9 a game in the Minnesota offense, at an ESPN "
                      "#94 price. That is within 11 points of quarterbacks going 50 picks "
                      "earlier — the reason waiting works."},
            {"player": "Brian Thomas Jr.", "badge": "IF YOU HAVE A QB", "kind": "alt",
             "blurb": "The best non-quarterback left on your board by value over "
                      "replacement. Take him if Lamar is already on your roster."},
            {"player": "Tyrone Tracy Jr.", "badge": "HANDCUFF UPSIDE", "kind": "alt",
             "blurb": "Unranked by ESPN entirely, model #71. The direct handcuff to "
                      "Skattebo — worth a bench spot if you took the Giants backfield."},
        ],
    },
    {
        "round": 10, "title": "Upside only — nothing here changes your season",
        "intro": "Everything left projects within a few points of replacement level. "
                 "Take the highest-upside names, then spend the last rounds on a defense "
                 "and a kicker.",
        "cards": [
            {"player": "Quentin Johnston", "badge": "THE PICK", "kind": "pick",
             "blurb": "168.7 points and 12.1 a game at an ESPN #127 price. The largest "
                      "remaining gap between your model and the consensus."},
            {"player": "Jacory Croskey-Merritt", "badge": "HANDCUFF UPSIDE", "kind": "alt",
             "blurb": "15 projected games in a Washington backfield with no settled "
                      "starter. Cheap lottery ticket."},
            {"player": "Dalton Schultz", "badge": "SAFE FILLER", "kind": "alt",
             "blurb": "Only if you still have no tight end. 16.1 projected games — the "
                      "most of any tight end on the board — but replacement-level points."},
        ],
    },
]


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--slot", type=int, default=5)
    ap.add_argument("--out", type=Path,
                    default=Path.home() / "Desktop" / "2026_Draft_Guide_Pick5.docx")
    args = ap.parse_args()
    build(args.slot, args.out, ROUNDS_SPEC)


if __name__ == "__main__":
    main()
